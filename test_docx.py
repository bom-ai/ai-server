#!/usr/bin/env python3
"""
실제 DOCX 파일로 테스트하는 스크립트
"""
import sys
import os

# 프로젝트 루트를 Python 경로에 추가
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from app.utils.docx_processor import (
    extract_text_with_separated_tables,
    extract_research_user_groups,
    extract_table_headers_with_subitems,
    format_items_for_prompt
)

def debug_docx_structure(file_path):
    """DOCX 파일의 상세 구조를 분석합니다."""
    try:
        from docx import Document
        import io
        
        print(f"=== DOCX 구조 분석: {file_path} ===")
        
        with open(file_path, 'rb') as f:
            file_content = f.read()
        
        file_stream = io.BytesIO(file_content)
        doc = Document(file_stream)
        
        print(f"총 문단 수: {len(doc.paragraphs)}")
        print(f"총 테이블 수: {len(doc.tables)}")
        
        print("\n--- 문단별 상세 정보 ---")
        for i, paragraph in enumerate(doc.paragraphs[:10]):  # 처음 10개만
            print(f"문단 {i+1}:")
            print(f"  스타일: {paragraph.style.name}")
            print(f"  텍스트: '{paragraph.text[:50]}{'...' if len(paragraph.text) > 50 else ''}'")
            print(f"  길이: {len(paragraph.text)}자")
            print()
        
        if len(doc.paragraphs) > 10:
            print(f"... 나머지 {len(doc.paragraphs) - 10}개 문단 생략")
        
        print("\n--- 테이블별 상세 정보 ---")
        for i, table in enumerate(doc.tables):
            print(f"테이블 {i+1}: {len(table.rows)}행 x {len(table.columns)}열")
            for row_idx, row in enumerate(table.rows[:3]):  # 처음 3행만
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()[:20] + ('...' if len(cell.text.strip()) > 20 else '')
                    row_text.append(cell_text)
                print(f"  행 {row_idx+1}: {' | '.join(row_text)}")
            if len(table.rows) > 3:
                print(f"  ... 나머지 {len(table.rows) - 3}행 생략")
            print()
            
    except Exception as e:
        print(f"구조 분석 중 오류: {e}")

def test_real_docx_file(file_path):
    """실제 DOCX 파일을 테스트합니다."""
    try:
        print(f"파일 테스트: {file_path}")
        
        # 파일 읽기
        with open(file_path, 'rb') as f:
            file_content = f.read()
        
        print(f"파일 크기: {len(file_content)} bytes")
      
        # 테이블 헤더만 추출 (custom_items 프롬프트 삽입용)  
        print("\n=== 분리된 테이블 추출 결과 ===")
        separated_data = extract_text_with_separated_tables(file_content)
        print(f"separated_data 키들: {list(separated_data.keys())}")
        print("-" * 50)
        
        print(f"paragraphs (개수: {len(separated_data['paragraphs'])}):")
        for i, para in enumerate(separated_data['paragraphs'][:3]):  # 처음 3개만
            print(f"  {i+1}: {para[:100]}{'...' if len(para) > 100 else ''}")
        if len(separated_data['paragraphs']) > 3:
            print(f"  ... 나머지 {len(separated_data['paragraphs']) - 3}개 문단")
        
        print(f"\ntable_headers (개수: {len(separated_data['table_headers'])}):")
        for header in separated_data['table_headers']:
            print(f"  테이블 {header['table_index']}: {header['content']}")
        
        print(f"\ntable_data_rows (개수: {len(separated_data['table_data_rows'])}):")
        for i, row in enumerate(separated_data['table_data_rows'][:5]):  # 처음 5개만
            print(f"  {i+1}. 테이블{row['table_index']}-행{row['row_index']}: {row['content'][:80]}{'...' if len(row['content']) > 80 else ''}")
        if len(separated_data['table_data_rows']) > 5:
            print(f"  ... 나머지 {len(separated_data['table_data_rows']) - 5}개 행")
        
        # 리서치 사용자 그룹 추출 메서드
        print(f"\n=== 리서치 사용자 그룹 추출 결과 ===")
        research_groups_data = extract_research_user_groups(separated_data['table_data_rows'])
        print(f"research_groups_data 키들: {list(research_groups_data.keys())}")
        print("-" * 50)
        
        print(f"총 고유 그룹 수: {research_groups_data['total_unique_groups']}")
        print(f"반복 등장한 그룹 수: {research_groups_data['total_repeated_groups']}")
        
        print(f"\nunique_groups (개수: {len(research_groups_data['unique_groups'])}):")
        for i, group in enumerate(research_groups_data['unique_groups']):
            print(f"  {i+1}. 테이블{group['table_index']}-행{group['row_index']}: {group['content'][:80]}{'...' if len(group['content']) > 80 else ''}")
        
        if research_groups_data['group_occurrence_stats']:
            print(f"\ngroup_occurrence_stats (그룹 등장 통계):")
            for group_info, count in list(research_groups_data['group_occurrence_stats'].items()):
                print(f"  '{group_info[:50]}{'...' if len(group_info) > 50 else ''}' → {count}회 등장")
        else:
            print(f"\ngroup_occurrence_stats: 중복 등장한 그룹이 없습니다.")
        
        # 새로운 기능: 테이블 헤더와 세부 항목 추출
        print(f"\n=== 테이블 헤더 및 세부 항목 추출 결과 ===")
        structured_items = extract_table_headers_with_subitems(file_content)
        print(f"구조화된 항목 수: {len(structured_items)}")
        print("-" * 50)
        
        for i, item in enumerate(structured_items):
            print(f"{i+1}. 헤더: {item['header']}")
            if item['subitems']:
                print(f"   세부항목: {item['subitems']}")
            else:
                print(f"   세부항목: 없음")
            print()
        
        # 프롬프트용 포맷팅 테스트
        print(f"=== 프롬프트용 포맷팅 결과 ===")
        formatted_items = format_items_for_prompt(structured_items)
        print(f"포맷팅된 항목 수: {len(formatted_items)}")
        print("-" * 50)
        
        print(formatted_items)
        
        print("-" * 50)
        
        print("✅ 테스트 완료!")
        
    except Exception as e:
        print(f"❌ 오류: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        if os.path.exists(file_path):
            # 구조 분석
            debug_docx_structure(file_path)
            print("\n" + "="*60 + "\n")
            # 함수 테스트
            test_real_docx_file(file_path)
        else:
            print(f"파일을 찾을 수 없습니다: {file_path}")
    else:
        print("사용법: python3 test_real_docx.py <docx_file_path>")
        print("예시: python3 test_real_docx.py ./sample.docx")
        print("\n추가 옵션:")
        print("  - 구조만 분석: python3 test_real_docx.py --debug <file>")
        
        # 간단한 옵션 처리
        if len(sys.argv) == 3 and sys.argv[1] == "--debug":
            file_path = sys.argv[2]
            if os.path.exists(file_path):
                debug_docx_structure(file_path)
            else:
                print(f"파일을 찾을 수 없습니다: {file_path}")
