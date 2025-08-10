"""
DOCX 파일 처리 유틸리티
"""
import io
import re
from io import BytesIO
from docx import Document
from typing import List, Dict, Optional, Any


def normalize_key(text: str) -> str:
    """
    텍스트를 정규화하여 키로 사용할 수 있도록 변환합니다.
    공백 제거, 소문자 변환, 특수문자 제거 등을 수행합니다.
    """
    if not text:
        return ""
    
    # 앞뒤 공백 제거 및 연속된 공백을 하나로 변환
    normalized = re.sub(r'\s+', ' ', text.strip())
    
    # 소문자로 변환
    normalized = normalized.lower()
    
    return normalized


def extract_text_with_separated_tables(file_content: bytes) -> dict:
    """
    DOCX 파일에서 텍스트를 추출하되, 테이블 헤더와 데이터를 분리합니다.
    
    Args:
        file_content: DOCX 파일의 바이트 내용
        
    Returns:
        {
            'paragraphs': 문단들의 리스트,
            'table_headers': 각 테이블 첫 번째 행들의 리스트,
            'table_data_rows': 각 테이블 나머지 행들의 리스트,
            'full_text': 전체 텍스트
        }
    """
    try:
        # 바이트 데이터를 메모리 스트림으로 변환
        file_stream = io.BytesIO(file_content)
        
        # Document 객체 생성
        doc = Document(file_stream)
        
        # 문단 추출
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 빈 줄 제외
                paragraphs.append(paragraph.text.strip())
        
        # 테이블에서 헤더와 데이터 분리 추출
        table_headers = []  # 각 테이블의 첫 번째 행(헤더)
        table_data_rows = []  # 각 테이블의 나머지 행들
        all_tables = []  # 전체 테이블 (기존 형태)
        
        for table_idx, table in enumerate(doc.tables):
            table_text = []
            
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                
                if row_text:
                    row_content = " | ".join(row_text)
                    table_text.append(row_content)
                    
                    # 첫 번째 행은 헤더로 분류
                    if row_idx == 0:
                        table_headers.append({
                            'table_index': table_idx,
                            'content': row_content
                        })
                    else:
                        # 나머지 행들은 데이터 행으로 분류
                        table_data_rows.append({
                            'table_index': table_idx,
                            'row_index': row_idx,
                            'content': row_content
                        })
            
            if table_text:
                all_tables.append("\n".join(table_text))
        
        # 전체 텍스트 결합
        all_text = []
        
        if paragraphs:
            all_text.append("=== 문서 내용 ===")
            all_text.extend(paragraphs)
        
        if all_tables:
            all_text.append("\n=== 표 내용 ===")
            all_text.extend(all_tables)
        
        return {
            'paragraphs': paragraphs,
            'table_headers': table_headers,
            'table_data_rows': table_data_rows,
            'full_text': "\n".join(all_text)
        }
        
    except Exception as e:
        raise Exception(f"DOCX 분리 처리 중 오류 발생: {str(e)}")


def extract_research_user_groups(table_data_rows: list) -> dict:
    """
    테이블 데이터에서 리서치 사용자 그룹을 추출하고 중복을 제거합니다.
    
    FGD 분석 템플릿에서 각 테이블은 동일한 사용자 그룹들에 대한 정보를 담고 있습니다.
    이 함수는 중복된 그룹 정보를 제거하여 고유한 리서치 대상 그룹들을 식별합니다.
    
    Args:
        table_data_rows: extract_text_with_separated_tables에서 반환된 table_data_rows
        
    Returns:
        {
            'unique_groups': 고유한 사용자 그룹들,
            'group_occurrence_stats': 각 그룹의 등장 횟수 통계,
            'total_unique_groups': 총 고유 그룹 수,
            'total_repeated_groups': 반복 등장한 그룹 수
        }
    """
    try:
        group_count = {}
        unique_groups = []
        
        for row_data in table_data_rows:
            group_info = row_data['content'].strip()  # 앞뒤 화이트스페이스 제거
            group_info = re.sub(r'\s+', ' ', group_info)  # 연속된 공백을 하나로 통합
            
            if group_info not in group_count:
                # 정리된 group_info로 새로운 row_data 생성
                cleaned_group_data = row_data.copy()
                cleaned_group_data['content'] = group_info
                
                group_count[group_info] = {
                    'count': 1,
                    'first_occurrence': cleaned_group_data
                }
                unique_groups.append(cleaned_group_data)
            else:
                group_count[group_info]['count'] += 1
        
        # 그룹 등장 횟수 통계 생성
        group_occurrence_stats = {}
        for group_info, info in group_count.items():
            if info['count'] > 1:
                group_occurrence_stats[group_info] = info['count']
        
        return {
            'unique_groups': unique_groups,
            'group_occurrence_stats': group_occurrence_stats,
            'total_unique_groups': len(unique_groups),
            'total_repeated_groups': len(group_occurrence_stats)
        }
        
    except Exception as e:
        raise Exception(f"리서치 사용자 그룹 추출 중 오류 발생: {str(e)}")


def extract_table_headers_with_subitems(file_content: bytes) -> List[Dict]:
    """
    DOCX 파일에서 테이블 헤더와 해당 테이블의 세부 항목들을 추출합니다. (디버깅 모드)
    """
    print("➡️ 함수 'extract_table_headers_with_subitems' 실행 시작")
    
    try:
        if not isinstance(file_content, bytes):
            print(f"❌ 오류: 입력된 file_content가 bytes 타입이 아닙니다. (타입: {type(file_content)})")
            raise TypeError("a bytes-like object is required, not 'str'")

        print(f"   - 입력된 파일 크기: {len(file_content)} bytes")
        file_stream = io.BytesIO(file_content)
        
        doc = Document(file_stream)
        print("   📄 DOCX 파일 로드 성공")
        
        structured_items = []
        print(f"   - 문서에서 총 {len(doc.tables)}개의 테이블 발견")
        
        for table_idx, table in enumerate(doc.tables):
            print(f"\n🔍 {table_idx}번 테이블 처리 중...")
            if len(table.rows) == 0:
                print("   - 테이블에 행이 없어 건너뜁니다.")
                continue
                
            header_row = table.rows[0]
            header_text = ""
            
            # 헤더 텍스트 후보들을 모두 확인
            header_candidates = [cell.text.strip() for cell in header_row.cells]
            print(f"   - 헤더 행 후보 텍스트: {header_candidates}")

            for cell_text in header_candidates:
                if cell_text:
                    header_text = cell_text
                    print(f"   - 테이블 헤더를 '{header_text}'로 확정")
                    break
            
            if not header_text:
                print("   - 유효한 헤더를 찾지 못해 건너뜁니다.")
                continue
                
            subitems = []
            
            for row_idx, row in enumerate(table.rows[1:], start=1):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if not cell_text:
                        continue
                    
                    # print(f"     - [행:{row_idx}, 열:{col_idx}] 셀 내용 확인: \"{cell_text[:30]}...\"")
                    lines = cell_text.split('\n')
                    for line in lines:
                        line = line.strip()
                        
                        if not line or len(line) < 2: continue
                        if re.match(r'^\d+-\d+$', line): continue
                        if len(line.split()) == 1 and len(line) < 10: continue
                        
                        item_text = None
                        if line.startswith('- '):
                            item_text = line[2:].strip()
                        elif line.startswith('• '):
                            item_text = line[2:].strip()
                        elif re.match(r'^\d+[\.\)]\s+', line):
                            item_text = re.sub(r'^\d+[\.\)]\s+', '', line).strip()
                        elif (len(line) > 10 and len(line) < 200 and not re.match(r'^\d+$', line) and '|' not in line):
                            item_text = line
                        
                        if item_text and item_text not in subitems and len(item_text) < 200:
                            print(f"    ✔️ ['{item_text}'] 항목 추가")
                            subitems.append(item_text)

            print(f"   - '{header_text}' 헤더에 총 {len(subitems)}개의 세부 항목 추출 완료.")
            structured_items.append({
                'header': header_text,
                'subitems': subitems,
                'table_index': table_idx
            })
        
        print(f"\n🏁 함수 실행 완료. 총 {len(structured_items)}개의 구조화된 항목 반환.")
        return structured_items
        
    except Exception as e:
        print(f"❌ 함수 실행 중 치명적인 오류 발생: {str(e)}")
        # 원래 오류를 포함하여 새로운 예외를 발생시켜, 어디서 문제가 생겼는지 추적하기 쉽게 함
        raise Exception(f"테이블 헤더 및 세부 항목 추출 중 오류 발생: {str(e)}")


def format_items_for_prompt(structured_items: List[Dict]) -> str:
    """
    구조화된 테이블 아이템들을 프롬프트용 계층적 문자열로 변환합니다.
    
    Args:
        structured_items: extract_table_headers_with_subitems에서 반환된 구조화된 아이템들
        
    Returns:
        프롬프트에 사용할 수 있는 계층적 구조의 문자열
        예:
        1. 크리에이터 현황
          - 활동 현황
          - 활동 목표
        2. 라이브 이용 현황
          - 라이브 현황
          - 가로/세로
    """
    try:
        formatted_lines = []
        
        for idx, item in enumerate(structured_items, 1):
            header = item['header']
            subitems = item['subitems']
            
            # 헤더를 번호와 함께 추가
            formatted_lines.append(f"{idx}. {header}")
            
            # 세부 항목들을 들여쓰기와 함께 추가
            if subitems:
                for subitem in subitems:
                    formatted_lines.append(f"  - {subitem}")
        
        return '\n'.join(formatted_lines)
        
    except Exception as e:
        raise Exception(f"프롬프트용 아이템 포맷팅 중 오류 발생: {str(e)}")

def parse_analysis_sections_any(analysis_text: str) -> Dict[str, str]:
    """
    'analysis' 문자열에서 번호 섹션(### 1. ...)별 본문을 추출해
    { '1. 제목': '본문', ... } 형태로 반환.
    """
    text = analysis_text.replace("\r\n", "\n").replace("\r", "\n")
    sec_pat = re.compile(r'(?m)^###\s*(\d+)\.\s*(.+?)\s*$')
    sections = list(sec_pat.finditer(text))

    out: Dict[str, str] = {}
    for i, m in enumerate(sections):
        num = m.group(1)
        title = m.group(2).strip()
        start = m.end()
        end = sections[i+1].start() if i+1 < len(sections) else len(text)
        body = text[start:end].strip()
        out[f"{num}. {title}"] = body
    return out

def replace_analysis_with_parsed(job_result: Dict[str, Any]) -> Dict[str, Any]:
    """
    job_result 내 모든 파일에 대해:
      - 'analysis'가 문자열이면 섹션 파싱하고
      - 파싱 성공 시 'analysis'를 딕셔너리로 교체
      - 파싱 실패/섹션 미존재 시 원문 유지
    최종적으로 수정된 job_result 반환.
    """
    results = job_result.get("results") or {}
    for file_key, file_obj in results.items():
        analysis_text = file_obj.get("analysis")
        if isinstance(analysis_text, str) and analysis_text.strip():
            parsed = parse_analysis_sections_any(analysis_text)
            if parsed:  # 섹션을 하나라도 찾았을 때만 교체
                file_obj["analysis"] = parsed
            # else: 섹션 패턴이 없으면 그대로 둠
    return job_result


def fill_frame_with_analysis_bytes(json_data: dict, frame_docx_bytes: bytes) -> bytes:
    """
    JSON 객체와 DOCX bytes를 받아, 분석 내용을 DOCX에 채워 넣고
    수정된 DOCX를 bytes로 반환
    """
    results = json_data.get("results", {})
    if not results:
        print("⚠️ JSON 안에 results가 없습니다.")
        return frame_docx_bytes

    # group -> {header -> 분석내용} 매핑 생성 (키 모두 normalize)
    group_to_analysis: Dict[str, Dict[str, str]] = {}
    for obj in results.values():
        group = normalize_key(obj.get("group"))
        analysis = obj.get("analysis")
        if isinstance(analysis, dict):
            norm_analysis = {normalize_key(k): (v or "") for k, v in analysis.items()}
            group_to_analysis[group] = norm_analysis

    if not group_to_analysis:
        print("⚠️ (group, analysis dict) 매핑이 비어있습니다.")
        return frame_docx_bytes

    # DOCX 로드 (bytes → Document)
    doc = Document(BytesIO(frame_docx_bytes))
    filled_count = 0

    # 표 순회
    for table in doc.tables:
        if not table.rows or len(table.columns) < 2:
            continue

        # 헤더(첫 행) 정규화
        headers = [normalize_key(cell.text) for cell in table.rows[0].cells]

        # 데이터 행
        for row in table.rows[1:]:
            group_name_norm = normalize_key(row.cells[0].text)
            if group_name_norm not in group_to_analysis:
                continue
            item_to_result = group_to_analysis[group_name_norm]

            # 각 열 채우기
            for col_idx in range(1, len(row.cells)):
                header_norm = headers[col_idx]
                if header_norm in item_to_result:
                    analysis_text = item_to_result[header_norm].strip()
                    if analysis_text:
                        cell = row.cells[col_idx]
                        if cell.text.strip():
                            cell.add_paragraph("")
                        p = cell.add_paragraph()
                        run = p.add_run("[분석]")
                        run.bold = True
                        cell.add_paragraph(analysis_text)
                        filled_count += 1

    # 수정된 DOCX → bytes 변환
    output_stream = BytesIO()
    doc.save(output_stream)
    return output_stream.getvalue()

